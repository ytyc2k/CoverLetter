from docx import Document
import csv
with open('skills.csv') as skills_csv:
    reader = csv.reader(skills_csv)
    rownum = 0
    for ln in reader:
        if rownum == 0:
            pass
        else:
            RR = [i for i in ln[5].split('\n')]
            docname='CoverLetter-'+ln[4]+'-'+ln[2]+'.docx'
            if RR[0] == '':
                break
            document = Document('demo.docx')
            for p in document.paragraphs:
                for t in p.runs:
                    t.text = t.text.replace('#position', ln[2])
                    t.text = t.text.replace('#website', ln[0])
            table = document.tables[0]
            # data = []
            for i in range(1,len(RR)+1):
                table.rows[i].cells[0].text=RR[i-1]
            document.save(docname)
        rownum = rownum + 1